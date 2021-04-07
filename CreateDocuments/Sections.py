import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches, Pt, RGBColor


def create_doc():
    doc = Document()
    return doc


document = create_doc()


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def save(isim='Word.docx'):
    document.save(isim)


def yaz_header_table():
    print('header')
    section = document.sections[0]
    header = section.header
    # delete_paragraph(header.paragraphs[0])
    table = header.add_table(4, 5, Inches(6))
    column = table.columns[0]
    column.cells[0].merge(column.cells[-1])
    table.style = 'Table Grid'


def makeRun(content, paragraph, font_name='Times New Roman', font_size=12, font_bold=False, font_italic=False,
            font_underline=False, color=RGBColor(0, 0, 0), align='left', style='Normal', add_break=False):
    paragraph.style = style
    run = paragraph.add_run(content)
    if add_break:
        run.add_break(WD_BREAK.LINE_CLEAR_LEFT)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    if align.lower() == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == 'right':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == 'justify':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def yazBaslik(line):
    a = line.split('%BA%')
    paragraph = document.add_paragraph()
    for ba in a:
        makeRun(ba, paragraph, font_size=15, font_bold=True, align='center', add_break=True)


def yazListe(line,ln = 0 ,list='List Bullet'):
    a = line.split('%LI%')
    for ba in a:
        paragraph = document.add_paragraph(style=list)
        if(ln>0):
            makeRun((f'{str(ln)}.\t{ba} '), paragraph, style=list)
        else:
            makeRun((ba), paragraph, style=list)



def yazHeading(line):
    a = line.split('%HE%')
    for ba in a:
        paragraph = document.add_heading(level=2)
        makeRun(ba, paragraph, style='Heading 2', font_bold=True)


def yazPara(line):
    paragraph = document.add_paragraph()
    b = line.split('%KA%')
    for bolum in range(len(b)):
        if bolum % 2 == 0:
            makeRun(b[bolum], paragraph, style='Normal')
        else:
            makeRun(b[bolum], paragraph, font_bold=True, style='Normal')


def funcSecond(filename):
    with open(filename, 'r', encoding='utf-8') as rf:
        ln = 0
        for line in rf:
            x = re.findall(r"\B%[a-zA-Z][a-zA-Z]\b%", line)
            if x:
                if x[0] == '%BA%':
                    ln = 0
                    yazBaslik(line[4:-1])
                elif x[0] == '%LB%':
                    ln = 0
                    yazListe(line[4:-1])
                elif x[0] == '%LN%':
                    ln +=1
                    if ln:
                        yazListe(line[4:-1], ln, 'List')
                elif x[0] == '%PA%':
                    ln = 0
                    yazPara(line[4:-1])
                elif x[0] == '%HE%':
                    ln = 0
                    yazHeading(line[4:-1])
