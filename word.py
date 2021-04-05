import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

document = Document()


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


section = document.sections[0]
header = section.header
delete_paragraph(header.paragraphs[0])
table = header.add_table(4, 5, Inches(6))
column = table.columns[0]
column.cells[0].merge(column.cells[-1])
table.style = 'Table Grid'


def writedocx(content='', font_name='Times New Roman', font_size=12, font_bold=False, font_italic=False,
              font_underline=False, color=RGBColor(0, 0, 0),
              before_spacing=0, after_spacing=10, line_spacing=1.5, keep_together=True, keep_with_next=False,
              page_break_before=False,
              widow_control=False, align='left', style='Normal', paragraph=document.add_paragraph()):
    run = paragraph.add_run(text=content)
    paragraph.style = style
    # print(f'Run Style: {run.style}')
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before_spacing)
    paragraph_format.space_after = Pt(after_spacing)
    paragraph_format.line_spacing = line_spacing
    paragraph_format.keep_together = keep_together
    paragraph_format.keep_with_next = keep_with_next
    paragraph_format.page_break_before = page_break_before
    paragraph_format.widow_control = widow_control
    if align.lower() == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == 'right':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == 'justify':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def makeRun(content, paragraph, font_name='Times New Roman', font_size=12, font_bold=False, font_italic=False,
            font_underline=False, color=RGBColor(0, 0, 0), align='left', style='Normal', add_break=False):
    paragraph.style = style
    run = paragraph.add_run(content)
    if add_break:
        run.add_break(WD_BREAK.LINE_CLEAR_LEFT)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    if align.lower() == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == 'right':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == 'justify':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def funcOne():
    with open('tag.txt', 'r', encoding='utf-8') as rf:
        paragraph = document.add_paragraph(style='Normal')
        for line in rf:
            x = re.findall(r"\B%[a-zA-Z][a-zA-Z]\b%", line)
            if x:
                if x[0] == '%BA%':
                    print(paragraph)
                    writedocx(content=line[4:], font_bold=True, font_size=17, align='center', paragraph=paragraph)
                    # paragraph = document.add_paragraph(text=line[4:], style='Normal')
                    print(paragraph)
                    print('Başlık')
                elif x[0] == '%LI%':
                    paragraph = document.add_paragraph(style='List')
                    print(paragraph)
                    writedocx(content=line[4:], font_size=12, style='List Bullet', paragraph=paragraph)

                    print(paragraph)
                    print('Liste')
                elif x[0] == '%PA%':
                    writedocx(content=line[4:], font_size=12, style='Normal', paragraph=paragraph, after_spacing=0)

                    # print(line)
                    print(paragraph)
                    print('Paragraph')


def yazBaslik(line):
    a = line.split('%BA%')
    paragraph = document.add_paragraph()
    for ba in a:
        makeRun(ba, paragraph, font_size=15, font_bold=True, align='center', add_break=True)


def yazListe(line):
    a = line.split('%LI%')
    for ba in a:
        paragraph = document.add_paragraph(style='List Bullet')
        makeRun(ba, paragraph, style='List Bullet')


def yazHeading(line):
    a = line.split('%HE%')
    for ba in a:
        paragraph = document.add_heading(level=2)
        makeRun(ba, paragraph, style='Heading 2', font_bold=True)
        print("header")


def yazPara(line):
    paragraph = document.add_paragraph()
    b = line.split('%KA%')
    for bolum in range(len(b)):
        if bolum % 2 == 0:
            makeRun(b[bolum], paragraph, style='Normal')
        else:
            makeRun(b[bolum], paragraph, font_bold=True, style='Normal')


def funcSecond():
    with open('test_ho1.txt ', 'r', encoding='utf-8') as rf:
        for line in rf:
            x = re.findall(r"\B%[a-zA-Z][a-zA-Z]\b%", line)
            if x:
                if x[0] == '%BA%':
                    yazBaslik(line[4:-1])
                elif x[0] == '%LI%':
                    yazListe(line[4:-1])
                elif x[0] == '%PA%':
                    yazPara(line[4:-1])
                elif x[0] == '%HE%':
                    yazHeading(line[4:-1])


funcSecond()

document.save('word.docx')
